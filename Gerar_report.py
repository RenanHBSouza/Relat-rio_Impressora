#Salvar log no servidor
#Enviar csv para o email desejado
#Executar código sendo possível escolher o arquivo que quero analisar
#Pegar somente a coluna "Message"
#Conseguir informação de usuário e numero de paginas de cada linha
#Armazenamento por lista, utilizando de uma para colaborador e outra para numero de paginas (mesmo indice)
#Após ler toda a planilha gerar um relatório com os dados

import pandas as pd
import tkinter as tk
from tkinter import filedialog
import os
from docx import Document
from docx.shared import Pt

VALOR_IMPRESSAO = 1.06
IMPRESSORA = 'Impressora 2 Colorida'
TOP_OFENSORES = int(input("Quantidade de top ofensores: "))

def selecionar_arquivo():
    root = tk.Tk()      #Cria uma janela oculta
    root.withdraw()
    arquivo = filedialog.askopenfilename(       #Abre a janela para seleção de arquivo com o tipo já determinado
        title="Selecione o arquivo desejado",
        filetypes=[("CSV files","*.csv")]
    )
    if os.path.isfile(arquivo):
        print("é arquivo")
    return arquivo


def busca_nome(nome, lista):        #Verifica se o nome buscado está na lista, retornando assim seu indice
    indice = -1

    for i in range(0, len(lista)):
        if nome == lista[i]:
            indice = i

    return indice


def ordenador_uso(users, paginas):      #Ordena as listas de maneira decrescente
    lista_nome = []
    lista_pag = []
    
    while len(paginas) != 0:
        maior = paginas[0]
        pos = 0
        for i in range(0,len(paginas)):
            if paginas[i] > maior:
                maior = paginas[i]
                pos = i
        lista_pag.append(maior)
        lista_nome.append(users[pos])
        del users[pos]          #Função para remover itens pelo indice
        del paginas[pos]

    return lista_nome, lista_pag


def filtro_planilha(caminho, busca):        #Função que filtra uma planilha a partir de uma busca
    if isinstance(caminho, str) and os.path.isfile(caminho):       
        sheet = pd.read_csv(caminho)
        filtro = sheet["Message"].str.contains(busca, na=False)     #Cria filtro para separar somente linhas que contenham a coluna desejada
        planilha_filtrada = sheet[filtro]                       #Aplica o filtro na planilha original
        planilha_filtrada = planilha_filtrada["Message"]        #Separa somente a coluna "Message"
    else:
        sheet = caminho
        filtro = sheet.str.contains(busca, na=False)     #Cria filtro para separar somente linhas que contenham a busca desejada
        planilha_filtrada = sheet[filtro]                       #Aplica o filtro na planilha original

    return planilha_filtrada


def leitor_info(caminho):       #Gera listas de usuários e paginas a partir da planilha filtrada

    planilha_filtrada = filtro_planilha(caminho, IMPRESSORA)
    num_registros = planilha_filtrada.shape[0]     #Variavel que armazena numero de linhas do csv
 
    users = []      #Lista de todos os usuarios
    num_paginas = []    #Lista de paginas impressas pelo usuario

    for z in range(num_registros):
        linha = planilha_filtrada.iloc[z]       #linha pega um registro específico da tabela
        linha = list(linha)     #Transforma essa string em um vetor de caracteres
        x = 0       
        saida_nome = []       #Cria lista de caracteres vazios
        saida_num = []

        while x < len(linha):       #executa o laço enquanto o indice não chega ao numero de caracteres da linha
            if linha[x] == 'b':     #Procura por b e y de forma consecutiva, buscando referencia para saber o usuario
                if linha[x+1] == 'y':
                    y = x+3
                    while linha[y] != ' ':
                        saida_nome.append(linha[y])     #Ao encontrar joga todos os proximos caracteres na variavel de saida
                        y += 1
                    break       #interrompe o loop pois já encontrou o usuário
            x += 1
        nome_str = ''.join(saida_nome)     #Transforma lista de caracteres em string novamente

        while x < len(linha):       #executa o laço enquanto o indice não chega ao numero de caracteres da linha
            if linha[x] == 'd':     #Procura por d e : de forma consecutiva, buscando referencia para numero de paginas
                if linha[x+1] == ':':
                    y = x+3
                    while linha[y] != '.':
                        saida_num.append(linha[y])     #Ao encontrar joga todos os proximos caracteres na variavel de saida
                        y += 1
                    break       #interrompe o loop pois já encontrou o numero de paginas
            x += 1
        num_str = ''.join(saida_num)        #Transforma lista de caracteres em string novamente

        indice = busca_nome(nome_str, users)      #Verifica se o nome já está na lista
        if indice == -1:
            users.append(nome_str)    #Insere o nome na lista de usuários
            num_paginas.append(int(num_str))        #Insere o numero de paginas no mesmo indice que o seu usuario
        else:
            num_paginas[indice] += int(num_str)     #Em caso de usuario já adicionado soma o numero de paginas no indice encontrado

    users, num_paginas = ordenador_uso(users,num_paginas)       #Ordena as listas em ordem decrescente
    
    return users, num_paginas, planilha_filtrada


def busca_documentos(planilha, user):        #Busca nome dos documentos impressos por usuário
    documentos = []
    planilha_filtrada = filtro_planilha(planilha, user)
    num_registros = planilha_filtrada.shape[0]

    for z in range(0, num_registros):
        linha = planilha_filtrada.iloc[z]       #linha pega um registro específico da tabela
        linha = list(linha)     #Transforma essa string em um vetor de caracteres
        x = 0       
        saida_doc = []       #Cria lista de caracteres vazios

        while x < len(linha):       #executa o laço enquanto o indice não chega ao numero de caracteres da linha
            if linha[x] == ',':     #Procura por ',' para começar a inserir o nome do documento na lista
                y = x+2
                while linha[y] != ' ' or linha[y+1] != 'o' or linha[y+2] != 'w' or linha[y+3] != 'n' or linha[y+4] != 'e' or linha[y+5] != 'd':
                    saida_doc.append(linha[y])     #Ao encontrar joga todos os proximos caracteres na variavel de saida
                    y += 1
                break       #interrompe o loop pois já encontrou o documento
            x += 1
        nome_str = ''.join(saida_doc)     #Transforma lista de caracteres em string novamente
        documentos.append(nome_str) 

    return documentos


def Sum_vet(lista):     #Faz o somatorio de valores de um vetor numerico
    val = 0
    for i in range(0,len(lista)):
        val += lista[i]

    return val


def gerar_relatorio(users, paginas, top, planilha):   #Gera o relatório txt a partir dos usuários e páginas listados
    num_pages =  Sum_vet(paginas)       #Armazena numero de paginas totais

    with open('report.txt', 'w') as file:   #Criação de um arquivo .txt

        #Cabeçalho do arquivo
        file.write("RELATÓRIO IMPRESSORA COLORIDA\n")
        file.write("==================================================\n\n")

        file.write("LISTA DE USO:\n\n")
        for x in range(0,len(users)):       #Inserção da lista de uso dos colaboradores
            linha = f"{users[x]:<30} {paginas[x]}"
            file.write(linha+"\n")
        file.write("==================================================\n\n")

        file.write("INFO. GERAIS DE IMPRESSORA:\n")       #Escreve informações da impressora
        file.write(f"\n{"Impressões Totais":<30} {num_pages}\n")
        file.write(f"{"Gasto Estimado":<30} {num_pages*VALOR_IMPRESSAO:.2f} reais\n")
        file.write("==================================================\n\n")

        file.write("INFO. PRINCIPAIS UTILIZADORES:\n")    #Informações mais especificas do colaborador com mais uso
        for i in range(0,top):
            file.write(f"\n{"Usuário":<30} {users[i]}\n")
            file.write(f"{"Gasto Estimado":<30} {paginas[i]*VALOR_IMPRESSAO:.2f} reais\n\n")
            docs = busca_documentos(planilha, users[i])
            for j in range(0,len(docs)):
                file.write(f"{docs[j]} \n")
        
        file.write("==================================================\n\n")

    return


def Relatorio_word(users, paginas, top, planilha):      #Gera um relatório Word a partir das informações coletadas
    num_pages = Sum_vet(paginas)

    doc = Document()        

    doc.add_heading('RELATÓRIO IMPRESSORA COLORIDA')        #Cabeçalho principal

    doc.add_heading('Lista de Uso', level=1)
    p1 = doc.add_paragraph()
    for x in range(0,len(users)):       #Inserção da lista de uso dos colaboradores
       p1.add_run('{}: '.format(users[x])).bold = True
       p1.add_run('{} página(s)\n'.format(paginas[x]))

    doc.add_heading('Informações da Impressora',level=1)        #informações acerca da impressora
    p2 = doc.add_paragraph()
    p2.add_run('Páginas Totais: ').bold = True
    p2.add_run(str(num_pages))
    p2.add_run('\nGasto estimado: ').bold = True
    valor = str("{:.2f}".format(num_pages*VALOR_IMPRESSAO))
    p2.add_run('R$ {}'.format(valor))

    doc.add_heading('Principais Utilizadores',level=1)      #Traz dados dos principais utilizadores definidos pelo administrador
    
    for i in range(0,top):
        p3 = doc.add_paragraph(style='List Number')
        p3.add_run('Usuário: ').bold = True
        p3.add_run(users[i])
        p3.add_run('\nGasto Estimado: ').bold = True
        valor = paginas[i]*VALOR_IMPRESSAO
        p3.add_run("R$ {:.2f}\n".format(valor))
        
        docs = busca_documentos(planilha, users[i])     #Busca os nomes dos documentos de cada user que está no top

        for j in range(0,len(docs)):
            p3.add_run(docs[j],style='Intense Quote Char').font.size = Pt(9)
            p3.add_run('\n')

    doc.save('Relatorio.docx')      #Salva o documento como "Relatorio.docx"
    return


path = selecionar_arquivo()
lista1, lista2, planilha = leitor_info(path)
Relatorio_word(lista1, lista2, TOP_OFENSORES, planilha)